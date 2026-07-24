#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
מוסיף רשומה חדשה ל-AGENT_REPORT.docx, לעמוד של הסוכן הנכון.
עמוד 1 = Claude, עמוד 2 = Antigravity. הרשומה החדשה תמיד בראש העמוד.

סימון ברור: שורות שמתחילות ב-"✔" מסומנות ירוק (בוצע), שורות שמתחילות ב-"✗"
מסומנות אדום (נכשל/דולג).

עמידות לקובץ פתוח: אם הוורד פתוח ונעול, הרשומה נשמרת לקובץ תור (.pending)
ותתווסף אוטומטית בריצה הבאה כשהקובץ יהיה פנוי.

שימוש:
    python3 append_report.py <summary.txt> <claude|ag>
"""
import sys, os, json, datetime
from docx import Document
from docx.shared import Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
DOCX = os.path.join(HERE, "AGENT_REPORT.docx")
PENDING = os.path.join(HERE, "AGENT_REPORT.pending.jsonl")

H_CLAUDE = "דוח Claude — מה נעשה"
H_AG = "דוח Antigravity — מה נעשה"
INTRO = "רשומה חדשה נוספת בכל ריצה. הרשומה הכי חדשה מופיעה מיד מתחת לשורה הזו."
GREEN = RGBColor(0x2E, 0x7D, 0x32)
RED = RGBColor(0xC6, 0x28, 0x28)


def build_doc():
    doc = Document()
    doc.add_heading(H_CLAUDE, level=0)
    doc.add_paragraph(INTRO)
    doc.add_page_break()
    doc.add_heading(H_AG, level=0)
    doc.add_paragraph(INTRO)
    doc.save(DOCX)


def get_anchor(doc, label):
    heading = H_CLAUDE if label == "Claude" else H_AG
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        if p.text.strip() == heading:
            return paras[i + 1]._p
    raise RuntimeError("anchor heading not found for " + label)


def insert_entry(doc, ts, label, body):
    anchor = get_anchor(doc, label)
    new = []
    head = doc.add_paragraph()
    r = head.add_run("■ " + ts + "  [" + label + "]")
    r.bold = True
    r.font.size = Pt(13)
    new.append(head._p)
    for line in body.splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line)
        s = line.strip()
        if s.startswith("✔"):
            run.font.color.rgb = GREEN
            run.bold = True
        elif s.startswith("✗"):
            run.font.color.rgb = RED
            run.bold = True
        new.append(p._p)
    new.append(doc.add_paragraph("─" * 30)._p)
    for p in reversed(new):
        anchor.addnext(p)


def load_pending():
    if not os.path.exists(PENDING):
        return []
    out = []
    with open(PENDING, encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if line:
                out.append(json.loads(line))
    return out


def save_pending(entries):
    if not entries:
        if os.path.exists(PENDING):
            os.remove(PENDING)
        return
    with open(PENDING, "w", encoding="utf-8") as f:
        for e in entries:
            f.write(json.dumps(e, ensure_ascii=False) + "\n")


def main():
    if len(sys.argv) < 3:
        print("usage: append_report.py <summary.txt> <claude|ag>", file=sys.stderr)
        sys.exit(1)
    with open(sys.argv[1], encoding="utf-8") as f:
        body = f.read().strip()
    a = sys.argv[2].strip().lower()
    label = "Claude" if a in ("claude", "c") else "Antigravity"
    ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")

    # התור: רשומות שממתינות מריצות קודמות + הרשומה החדשה
    entries = load_pending()
    entries.append({"ts": ts, "label": label, "body": body})

    if not os.path.exists(DOCX):
        try:
            build_doc()
        except PermissionError:
            save_pending(entries)
            print("Word file busy — entry queued to pending.")
            return

    try:
        doc = Document(DOCX)
        for e in entries:
            insert_entry(doc, e["ts"], e["label"], e["body"])
        doc.save(DOCX)
    except PermissionError:
        save_pending(entries)
        print("Word file is OPEN/locked — " + str(len(entries)) +
              " entry(ies) queued to .pending, will be added next run.")
        return

    save_pending([])  # הצליח — נקה את התור
    print("Report saved to " + label + " page (" + str(len(entries)) +
          " entry/ies flushed): " + ts)


if __name__ == "__main__":
    main()
